"""
Export Routes - v29.0
Eksport artykułów do DOCX/HTML/TXT + Editorial Review (Claude API)

ZMIANY v29.0:
- DIFF-BASED: Claude zwraca tylko zmiany, nie przepisuje całości
- BURSTINESS VALIDATION: Sprawdza CV przed/po
- AUTO-ROLLBACK: Automatycznie przywraca oryginał jeśli CV spadło
- DIFF OUTPUT: Szczegółowa lista zmian applied/failed
- MANUAL ROLLBACK: Nowy endpoint /editorial_review/rollback

ZMIANY v27.1:
- rescan_keywords_after_editorial() - przelicza frazy po Claude review

ZMIANY v27.0:
- Zmiana z Gemini na Claude API
- Wplata nieużyte frazy BASIC/EXTENDED
"""

import os
import re
import io
import json
import statistics
from datetime import datetime
from flask import Blueprint, request, jsonify, Response

# Firebase Firestore only
from firebase_admin import firestore

# Document generation
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Claude API for Editorial Review
import anthropic

# v27.1: Import keyword counter do rescan
try:
    from keyword_counter import count_keywords, count_keywords_for_state
    KEYWORD_COUNTER_OK = True
    print("[EXPORT] ✅ keyword_counter imported for rescan")
except ImportError:
    KEYWORD_COUNTER_OK = False
    print("[EXPORT] ⚠️ keyword_counter not available - rescan will use fallback")

# v29.0: Burstiness validation for editorial review
try:
    from ai_detection_metrics import calculate_burstiness
    BURSTINESS_CHECK_OK = True
    print("[EXPORT] ✅ Burstiness validation enabled")
except ImportError:
    BURSTINESS_CHECK_OK = False
    print("[EXPORT] ⚠️ Burstiness validation not available")

# v27.1: Ładowanie promptu z pliku (opcjonalne)
EDITORIAL_PROMPT_TEMPLATE = None
try:
    import os
    prompt_path = os.path.join(os.path.dirname(__file__), "editorial_prompt.json")
    if os.path.exists(prompt_path):
        with open(prompt_path, "r", encoding="utf-8") as f:
            prompt_data = json.load(f)
            EDITORIAL_PROMPT_TEMPLATE = prompt_data.get("full_prompt_assembled")
            print(f"[EXPORT] ✅ Loaded editorial prompt from file (v{prompt_data.get('version', '?')})")
except Exception as e:
    print(f"[EXPORT] ℹ️ Using built-in prompt (file load failed: {e})")

export_routes = Blueprint("export_routes", __name__)

# v51: Safe full_article extraction — handles both string and dict formats
def _safe_get_full_article(project_data):
    """Extract text from full_article field, handling both dict and string."""
    fa = project_data.get("full_article")
    if isinstance(fa, dict):
        return fa.get("content", "")
    elif isinstance(fa, str):
        return fa
    return ""

# Claude config
ANTHROPIC_API_KEY = os.getenv("ANTHROPIC_API_KEY")
# v50.5: Configurable editorial model via env var
# v51: Changed default to alias (no snapshot) + added fallback chain
# Fix #29: Walidacja modelu — claude-sonnet-4-6-20250514 nie istnieje
_raw_model = os.getenv("EDITORIAL_MODEL", "claude-sonnet-4-5-20250929")
# Walidacja: jesli model zawiera snapshot date, sprawdz czy to znany model
_KNOWN_MODELS = {
    "claude-sonnet-4-5-20250929", "claude-haiku-4-5-20251001",
    "claude-opus-4-5-20250101",
    # aliasy bez daty tez OK
    "claude-sonnet-4-5", "claude-haiku-4-5", "claude-opus-4-5",
}
# Fix #29 v2: Wzmocniona walidacja — wylapuje literowki (np. "laude-" zamiast "claude-")
if not _raw_model.startswith("claude-"):
    print(f"[EXPORT] ⚠️ EDITORIAL_MODEL '{_raw_model}' nie zaczyna sie od 'claude-' (literowka?), falling back to claude-sonnet-4-5-20250929")
    _raw_model = "claude-sonnet-4-5-20250929"
elif _raw_model not in _KNOWN_MODELS and "-2025" in _raw_model:
    print(f"[EXPORT] ⚠️ EDITORIAL_MODEL '{_raw_model}' may not exist, falling back to claude-sonnet-4-5-20250929")
    _raw_model = "claude-sonnet-4-5-20250929"
EDITORIAL_MODEL = _raw_model
EDITORIAL_MODEL_FALLBACK = os.getenv("EDITORIAL_MODEL_FALLBACK", "claude-haiku-4-5-20251001")
claude_client = None
if ANTHROPIC_API_KEY:
    claude_client = anthropic.Anthropic(api_key=ANTHROPIC_API_KEY)
    print(f"[EXPORT] ✅ Claude API configured (editorial model: {EDITORIAL_MODEL}, fallback: {EDITORIAL_MODEL_FALLBACK})")
else:
    print("[EXPORT] ⚠️ ANTHROPIC_API_KEY not set")


def _robust_json_parse(text: str) -> dict | None:
    """
    Robust JSON parser for LLM responses.
    Handles: markdown fences, truncated corrected_text, raw newlines in strings.
    """
    import json, re

    # 1. Strip markdown fences
    clean = text.strip()
    clean = re.sub(r'^```json\s*', '', clean)
    clean = re.sub(r'^```\s*', '', clean)
    clean = re.sub(r'```\s*$', '', clean).strip()

    # 2. Find outer braces
    first = clean.find('{')
    last = clean.rfind('}')
    if first == -1 or last <= first:
        return None
    candidate = clean[first:last + 1]

    # 3. Try direct parse
    try:
        return json.loads(candidate)
    except Exception:
        pass

    # 4. Try with strict=False (allows control chars)
    try:
        return json.loads(candidate, strict=False)
    except Exception:
        pass

    # 5. Strip corrected_text field (often causes parse errors due to newlines)
    try:
        stripped = re.sub(
            r'"corrected_text"\s*:\s*"[\s\S]*?"(?=\s*[,}])',
            '"corrected_text": ""',
            candidate
        )
        result = json.loads(stripped, strict=False)
        result["corrected_text"] = ""  # mark as stripped
        return result
    except Exception:
        pass

    # 6. Extract just the safe fields
    try:
        score_m = re.search(r'"overall_score"\s*:\s*(\d+)', candidate)
        summary_m = re.search(r'"summary"\s*:\s*"([^"]{0,300})', candidate)
        errors_m = re.search(r'"errors_to_fix"\s*:\s*(\[[\s\S]*?\])', candidate)

        partial = {"overall_score": 5, "summary": "Analiza czesciowo sparsowana", "errors_to_fix": []}
        if score_m:
            partial["overall_score"] = int(score_m.group(1))
        if summary_m:
            partial["summary"] = summary_m.group(1)
        if errors_m:
            try:
                partial["errors_to_fix"] = json.loads(errors_m.group(1), strict=False)
            except Exception:
                pass
        return partial
    except Exception:
        pass

    return None



def _claude_call(prompt: str, max_tokens: int = 6000, label: str = "CALL") -> tuple:
    """Call Claude with model fallback. Returns (response_text, model_used)."""
    models_to_try = [EDITORIAL_MODEL, EDITORIAL_MODEL_FALLBACK]
    for model in models_to_try:
        try:
            print(f"[EDITORIAL_REVIEW] {label}: trying model={model}")
            resp = claude_client.messages.create(
                model=model,
                max_tokens=max_tokens,
                messages=[{"role": "user", "content": prompt}]
            )
            text = resp.content[0].text.strip()
            print(f"[EDITORIAL_REVIEW] {label}: ✅ model={model} → {len(text)} chars")
            return text, model
        except Exception as e:
            err_str = str(e)
            print(f"[EDITORIAL_REVIEW] {label}: ❌ model={model} → {err_str[:200]}")
            if model == models_to_try[-1]:
                raise  # re-raise if last model also fails
    raise RuntimeError("All Claude models failed")

# Fallback: Gemini (jeśli Claude niedostępny)
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")
genai = None
if GEMINI_API_KEY:
    try:
        import google.generativeai as genai
        genai.configure(api_key=GEMINI_API_KEY)
        if not ANTHROPIC_API_KEY:
            print("[EXPORT] ℹ️ Using Gemini as fallback")
    except ImportError:
        genai = None


def html_to_text(html: str) -> str:
    """Konwertuje HTML na czysty tekst."""
    text = re.sub(r'<h2[^>]*>(.*?)</h2>', r'\n\n## \1\n\n', html, flags=re.IGNORECASE | re.DOTALL)
    text = re.sub(r'<h3[^>]*>(.*?)</h3>', r'\n\n### \1\n\n', text, flags=re.IGNORECASE | re.DOTALL)
    text = re.sub(r'<p[^>]*>(.*?)</p>', r'\1\n\n', text, flags=re.IGNORECASE | re.DOTALL)
    text = re.sub(r'<br\s*/?>', '\n', text, flags=re.IGNORECASE)
    text = re.sub(r'<li[^>]*>(.*?)</li>', r'• \1\n', text, flags=re.IGNORECASE | re.DOTALL)
    text = re.sub(r'<[^>]+>', '', text)
    text = re.sub(r'\n{3,}', '\n\n', text)
    return text.strip()


def parse_article_structure(text: str) -> list:
    """Parsuje artykuł na strukturę nagłówków i paragrafów."""
    elements = []
    
    # Normalizuj format
    text = re.sub(r'^h2:\s*(.+)$', r'<h2>\1</h2>', text, flags=re.MULTILINE)
    text = re.sub(r'^h3:\s*(.+)$', r'<h3>\1</h3>', text, flags=re.MULTILINE)
    
    parts = re.split(r'(<h[23][^>]*>.*?</h[23]>)', text, flags=re.IGNORECASE | re.DOTALL)
    
    for part in parts:
        part = part.strip()
        if not part:
            continue
            
        h2_match = re.match(r'<h2[^>]*>(.*?)</h2>', part, re.IGNORECASE | re.DOTALL)
        h3_match = re.match(r'<h3[^>]*>(.*?)</h3>', part, re.IGNORECASE | re.DOTALL)
        
        if h2_match:
            elements.append({"type": "h2", "content": h2_match.group(1).strip()})
        elif h3_match:
            elements.append({"type": "h3", "content": h3_match.group(1).strip()})
        else:
            clean_text = html_to_text(part)
            if clean_text:
                elements.append({"type": "paragraph", "content": clean_text})
    
    return elements


# ================================================================
# v27.0: HELPER - Składanie treści z batchów
# ================================================================
def extract_text_from_batches(batches: list) -> tuple:
    """
    Bezpieczne składanie treści z batchów.
    Obsługuje różne formaty: batch.text, batch.content, batch.html, lub string
    
    Returns:
        (full_text, debug_info)
    """
    texts = []
    debug_info = {
        "total_batches": len(batches),
        "batches_with_text": 0,
        "batches_empty": 0,
        "fields_found": set()
    }
    
    for i, batch in enumerate(batches):
        text = None
        
        # v51: Handle string batches
        if isinstance(batch, str):
            text = batch.strip()
            if text:
                debug_info["fields_found"].add("string")
        elif isinstance(batch, dict):
            for field in ["text", "content", "html", "batch_text", "raw_text"]:
                if field in batch and batch[field]:
                    text = str(batch[field]).strip()
                    debug_info["fields_found"].add(field)
                    break
        
        if text:
            texts.append(text)
            debug_info["batches_with_text"] += 1
        else:
            debug_info["batches_empty"] += 1
            batch_keys = list(batch.keys()) if isinstance(batch, dict) else [type(batch).__name__]
            print(f"[EXPORT] ⚠️ Batch {i} jest pusty. Klucze: {batch_keys}")
    
    debug_info["fields_found"] = list(debug_info["fields_found"])
    full_text = "\n\n".join(texts)
    
    return full_text, debug_info


# ================================================================
# v27.1: RESCAN KEYWORDS AFTER EDITORIAL
# ================================================================
def rescan_keywords_after_editorial(project_id: str, corrected_article: str) -> dict:
    """
    v27.1: Przelicza frazy od ZERA po edycji Claude.
    """
    db = firestore.client()
    doc = db.collection("seo_projects").document(project_id).get()
    
    if not doc.exists:
        return {"error": "Project not found", "rescanned": False}
    
    project_data = doc.to_dict()
    keywords_state = project_data.get("keywords_state", {})
    
    if not keywords_state:
        return {"error": "No keywords_state", "rescanned": False}
    
    if not corrected_article or len(corrected_article.strip()) < 50:
        return {"error": "No corrected_article to scan", "rescanned": False}
    
    # Zbierz wszystkie keywordy
    keywords = []
    rid_to_keyword = {}
    for rid, meta in keywords_state.items():
        kw = (meta.get("keyword") or "").strip()
        if kw:
            keywords.append(kw)
            rid_to_keyword[rid] = kw
    
    if not keywords:
        return {"error": "No keywords to scan", "rescanned": False}
    
    # Przelicz frazy
    changes = []
    new_counts = {}
    
    if KEYWORD_COUNTER_OK:
        new_counts = count_keywords_for_state(corrected_article, keywords_state, use_exclusive_for_nested=False)
        print(f"[RESCAN] Using keyword_counter for {len(keywords)} keywords")
    else:
        # Fallback - proste liczenie
        clean_text = re.sub(r'<[^>]+>', ' ', corrected_article.lower())
        clean_text = re.sub(r'\s+', ' ', clean_text)
        
        for rid, kw in rid_to_keyword.items():
            kw_lower = kw.lower()
            count = clean_text.count(kw_lower)
            new_counts[rid] = count
        print(f"[RESCAN] Using fallback counter for {len(keywords)} keywords")
    
    # Porównaj i zaktualizuj
    for rid, meta in keywords_state.items():
        old_actual = meta.get("actual_uses", 0)
        new_actual = new_counts.get(rid, 0)
        
        if old_actual != new_actual:
            kw = meta.get("keyword", "?")
            kw_type = meta.get("type", "BASIC")
            changes.append({
                "keyword": kw,
                "type": kw_type,
                "old": old_actual,
                "new": new_actual,
                "diff": new_actual - old_actual
            })
            print(f"[RESCAN] '{kw}' ({kw_type}): {old_actual} → {new_actual}")
        
        # NADPISZ actual_uses
        meta["actual_uses"] = new_actual
        
        # Przelicz status
        min_t = meta.get("target_min", 1 if meta.get("type", "").upper() == "BASIC" else 1)
        max_t = meta.get("target_max", 999)
        
        if new_actual < min_t:
            meta["status"] = "UNDER"
        elif new_actual == max_t:
            meta["status"] = "OPTIMAL"
        elif min_t <= new_actual < max_t:
            meta["status"] = "OK"
        else:
            meta["status"] = "OVER"
        
        meta["remaining_max"] = max(0, max_t - new_actual)
        keywords_state[rid] = meta
    
    # Zapisz do Firebase
    try:
        db.collection("seo_projects").document(project_id).update({
            "keywords_state": keywords_state,
            "last_rescan": {
                "timestamp": firestore.SERVER_TIMESTAMP,
                "changes_count": len(changes),
                "trigger": "editorial_review"
            }
        })
        print(f"[RESCAN] ✅ Saved {len(changes)} changes to Firebase")
    except Exception as e:
        print(f"[RESCAN] ❌ Firebase save error: {e}")
        return {"error": str(e), "rescanned": False, "changes": changes}
    
    # Policz nowe coverage
    basic_total = 0
    basic_covered = 0
    extended_total = 0
    extended_covered = 0
    
    for rid, meta in keywords_state.items():
        kw_type = meta.get("type", "BASIC").upper()
        actual = meta.get("actual_uses", 0)
        
        if kw_type == "BASIC" or kw_type == "MAIN":
            basic_total += 1
            if actual >= 1:
                basic_covered += 1
        elif kw_type == "EXTENDED":
            extended_total += 1
            if actual >= 1:
                extended_covered += 1
    
    basic_pct = round(100 * basic_covered / basic_total, 1) if basic_total > 0 else 100
    extended_pct = round(100 * extended_covered / extended_total, 1) if extended_total > 0 else 100
    
    return {
        "rescanned": True,
        "changes_count": len(changes),
        "changes": changes,
        "coverage_after_rescan": {
            "basic": f"{basic_covered}/{basic_total} ({basic_pct}%)",
            "extended": f"{extended_covered}/{extended_total} ({extended_pct}%)"
        },
        "missing_after_rescan": {
            "basic": [c["keyword"] for c in changes if c["type"].upper() in ["BASIC", "MAIN", "ENTITY"] and c["new"] == 0],
            "extended": [c["keyword"] for c in changes if c["type"].upper() == "EXTENDED" and c["new"] == 0]
        }
    }


# ================================================================
# v29.0: DIFF PARSING HELPERS
# ================================================================
def parse_diff_response(response_text: str) -> list:
    """
    Parsuje odpowiedź Claude w formacie DIFF.
    
    Format wejściowy:
    [ZMIANA 1]
    ZNAJDŹ: "dokładny cytat"
    ZAMIEŃ: "poprawiona wersja"
    POWÓD: wyjaśnienie
    
    Zwraca: [{"find": "...", "replace": "...", "reason": "..."}]
    """
    changes = []
    
    # Wzorzec dla pojedynczej zmiany
    pattern = r'\[ZMIANA \d+\]\s*\n\s*ZNAJDŹ:\s*["\'](.+?)["\']\s*\n\s*ZAMIEŃ:\s*["\'](.+?)["\']\s*\n\s*POWÓD:\s*(.+?)(?=\n\s*\[ZMIANA|\Z)'
    
    matches = re.findall(pattern, response_text, re.DOTALL | re.IGNORECASE)
    
    for match in matches:
        find_text = match[0].strip()
        replace_text = match[1].strip()
        reason = match[2].strip()
        
        if find_text and len(find_text) > 5:
            changes.append({
                "find": find_text,
                "replace": replace_text,
                "reason": reason,
                "applied": False
            })
    
    # Fallback: prostszy wzorzec
    if not changes:
        simple_pattern = r'ZNAJDŹ:\s*["\']?(.+?)["\']?\s*\n\s*ZAMIEŃ:\s*["\']?(.+?)["\']?\s*(?:\n|$)'
        simple_matches = re.findall(simple_pattern, response_text, re.DOTALL)
        
        for match in simple_matches:
            find_text = match[0].strip().strip('"\'')
            replace_text = match[1].strip().strip('"\'')
            
            if find_text and len(find_text) > 5:
                changes.append({
                    "find": find_text,
                    "replace": replace_text,
                    "reason": "auto-parsed",
                    "applied": False
                })
    
    return changes


def _normalize_quotes(text: str) -> str:
    """
    v53: Normalizuje cudzysłowy do postaci ASCII dla porównań.
    Claude często zwraca „polskie" cudzysłowy, a tekst ma ASCII lub odwrotnie.
    """
    return (text
            .replace('\u201e', '"')   # „ → "
            .replace('\u201d', '"')   # " → "
            .replace('\u201c', '"')   # " → "
            .replace('\u00ab', '"')   # « → "
            .replace('\u00bb', '"')   # » → "
            .replace('\u2018', "'")   # ' → '
            .replace('\u2019', "'")   # ' → '
            .replace('\u2013', '-')   # – → -
            .replace('\u2014', '-')   # — → -
            )


def apply_diffs(original_text: str, changes: list) -> tuple:
    """
    Aplikuje zmiany diff do tekstu.
    v53: Dodano normalizację cudzysłowów (Fix 2).

    Zwraca: (modified_text, applied_changes, failed_changes)
    """
    modified = original_text
    applied = []
    failed = []

    for change in changes:
        find_text = change["find"]
        replace_text = change["replace"]

        # Próba 1: Dokładne dopasowanie
        if find_text in modified:
            modified = modified.replace(find_text, replace_text, 1)
            change["applied"] = True
            applied.append(change)
            continue

        # Próba 2: Dopasowanie z normalizacją cudzysłowów
        find_norm_quotes = _normalize_quotes(find_text)
        modified_norm_quotes = _normalize_quotes(modified)
        if find_norm_quotes in modified_norm_quotes:
            # Znajdź pozycję w znormalizowanym tekście i zamień w oryginale
            idx = modified_norm_quotes.index(find_norm_quotes)
            # Oblicz oryginalny fragment (ta sama pozycja, ta sama długość)
            original_fragment = modified[idx:idx + len(find_text)]
            # Sprawdź czy po normalizacji pasuje (zabezpieczenie)
            if _normalize_quotes(original_fragment) == find_norm_quotes:
                modified = modified[:idx] + replace_text + modified[idx + len(original_fragment):]
            else:
                # Fallback: szukaj po znormalizowanym i zamień
                modified = modified[:idx] + replace_text + modified[idx + len(find_norm_quotes):]
            change["applied"] = True
            change["match_type"] = "quote_normalized"
            applied.append(change)
            continue

        # Próba 3: Dopasowanie ignorujące whitespace + cudzysłowy
        find_normalized = _normalize_quotes(' '.join(find_text.split()))
        text_lines = modified.split('\n')
        found = False

        for i, line in enumerate(text_lines):
            line_normalized = _normalize_quotes(' '.join(line.split()))
            if find_normalized in line_normalized:
                # Zamień w tej linii (użyj oryginalnego find_text.strip())
                # Ale też spróbuj z quote-normalized
                if find_text.strip() in line:
                    text_lines[i] = line.replace(find_text.strip(), replace_text.strip(), 1)
                else:
                    # Zamień przez normalizację
                    line_norm = _normalize_quotes(line)
                    find_stripped_norm = _normalize_quotes(find_text.strip())
                    if find_stripped_norm in line_norm:
                        idx = line_norm.index(find_stripped_norm)
                        text_lines[i] = line[:idx] + replace_text.strip() + line[idx + len(find_stripped_norm):]
                    else:
                        text_lines[i] = line.replace(find_text.strip(), replace_text.strip(), 1)
                modified = '\n'.join(text_lines)
                change["applied"] = True
                change["match_type"] = "whitespace_quote_normalized"
                applied.append(change)
                found = True
                break

        if found:
            continue

        # Nie udało się
        change["applied"] = False
        change["error"] = "Fragment nie znaleziony w tekście"
        failed.append(change)

    return modified, applied, failed


def calculate_diff_stats(original: str, modified: str) -> dict:
    """Oblicza statystyki różnic między tekstami."""
    original_words = original.split()
    modified_words = modified.split()
    
    original_sentences = [s for s in re.split(r'[.!?]+', original) if s.strip()]
    modified_sentences = [s for s in re.split(r'[.!?]+', modified) if s.strip()]
    
    words_added = len(modified_words) - len(original_words)
    sentences_diff = len(modified_sentences) - len(original_sentences)
    
    # Procent zmian
    if len(original) > 0:
        changes = sum(1 for a, b in zip(original, modified) if a != b)
        changes += abs(len(original) - len(modified))
        change_percent = round((changes / len(original)) * 100, 1)
    else:
        change_percent = 100
    
    return {
        "words_original": len(original_words),
        "words_modified": len(modified_words),
        "words_diff": words_added,
        "sentences_original": len(original_sentences),
        "sentences_modified": len(modified_sentences),
        "sentences_diff": sentences_diff,
        "change_percent": min(change_percent, 100)
    }


# ================================================================
# EDITORIAL REVIEW - v29.0 (DIFF-BASED + BURSTINESS VALIDATION)
# ================================================================
@export_routes.post("/api/project/<project_id>/editorial_review")
def editorial_review(project_id):
    """
    v29.0: DIFF-BASED Editorial Review z walidacją burstiness.
    
    NOWOŚCI:
    - DIFF-BASED: Claude zwraca tylko zmiany, nie przepisuje całości
    - BURSTINESS VALIDATION: Sprawdza CV przed/po
    - AUTO-ROLLBACK: Automatycznie przywraca oryginał jeśli CV spadło
    - DIFF OUTPUT: Szczegółowa lista zmian
    """
    db = firestore.client()
    doc = db.collection("seo_projects").document(project_id).get()
    
    if not doc.exists:
        return jsonify({"error": "Project not found"}), 404
    
    project_data = doc.to_dict()
    
    # Pobierz treść
    full_text = None
    debug_info = {}
    
    # v51: Defensive — full_article can be string or dict
    fa_text = _safe_get_full_article(project_data)
    if fa_text:
        full_text = fa_text
        debug_info = {"source": "full_article"}
    else:
        batches = project_data.get("batches", [])
        if batches:
            full_text, debug_info = extract_text_from_batches(batches)
            debug_info["source"] = "batches"

    # v55.1: Word count guard — if full_article is truncated vs batches, use batches
    batches = project_data.get("batches", [])
    if full_text and batches:
        batches_text_parts = []
        for b in batches:
            if isinstance(b, dict):
                batches_text_parts.append(b.get("text", b.get("content", "")))
            elif isinstance(b, str):
                batches_text_parts.append(b)
        batches_wc = sum(len(p.split()) for p in batches_text_parts if p)
        full_text_wc = len(full_text.split())
        if batches_wc > 0 and full_text_wc < batches_wc * 0.7:
            print(f"[EDITORIAL_REVIEW] ⚠️ full_article truncated ({full_text_wc} vs {batches_wc} from batches) — using batches")
            full_text = "\n\n".join(p for p in batches_text_parts if p)
            debug_info = {"source": "batches_fallback", "reason": f"full_article truncated ({full_text_wc}<{batches_wc})"}

    if not full_text:
        return jsonify({
            "error": "No content found",
            "hint": "Użyj save_full_article lub addBatch aby zapisać treść.",
            "project_keys": list(project_data.keys())
        }), 400
    
    # Sprawdź API
    if not claude_client and not genai:
        return jsonify({"error": "No AI API configured (need ANTHROPIC_API_KEY or GEMINI_API_KEY)"}), 500
    
    word_count = len(full_text.split()) if full_text else 0
    
    if word_count < 50:
        return jsonify({
            "error": f"Article too short for review ({word_count} words)",
            "debug": debug_info,
            "hint": "Użyj save_full_article lub addBatch aby zapisać treść."
        }), 400
    
    topic = project_data.get("topic") or project_data.get("main_keyword", "artykuł")
    
    # ================================================================
    # v29.0: BURSTINESS CHECK - PRZED
    # ================================================================
    burstiness_before = None
    if BURSTINESS_CHECK_OK:
        try:
            burstiness_before = calculate_burstiness(full_text)
            print(f"[EDITORIAL_REVIEW] 📊 Burstiness PRZED: CV={burstiness_before.get('cv', 0)}")
        except Exception as e:
            print(f"[EDITORIAL_REVIEW] ⚠️ Burstiness check failed: {e}")
    
    # v27.0: Pobierz nieużyte frazy BASIC i EXTENDED
    keywords_state = project_data.get("keywords_state", {})
    unused_basic = []
    unused_extended = []
    
    for rid, meta in keywords_state.items():
        kw = meta.get("keyword", "")
        kw_type = meta.get("type", "BASIC").upper()
        actual = meta.get("actual_uses", 0)
        
        if actual == 0 and kw:
            if kw_type == "BASIC":
                unused_basic.append(kw)
            elif kw_type == "EXTENDED":
                unused_extended.append(kw)
    
    # Sekcja nieużytych fraz dla promptu
    unused_keywords_section = ""
    if unused_basic or unused_extended:
        unused_keywords_section = "\n=== ⚠️ NIEWYKORZYSTANE FRAZY SEO ===\n"
        if unused_basic:
            unused_keywords_section += f"BASIC (MUSZĄ być w tekście): {', '.join(unused_basic[:15])}\n"
        if unused_extended:
            unused_keywords_section += f"EXTENDED (wpleć jeśli pasują): {', '.join(unused_extended[:15])}\n"
    
    # v50.5: Overused keywords (stuffing) from Firebase
    overused_keywords = []
    for rid, meta in keywords_state.items():
        kw = meta.get("keyword", "")
        kw_type = meta.get("type", "BASIC").upper()
        actual = meta.get("actual_uses", 0)
        target_max = meta.get("target_max", 999)
        if actual > target_max and kw:
            overused_keywords.append({
                "keyword": kw,
                "actual": actual,
                "max": target_max,
                "excess": actual - target_max
            })
    
    if overused_keywords:
        overused_section = "\n=== 🔴 FRAZY NADUŻYTE (STUFFING — zmniejsz!) ===\n"
        for ow in sorted(overused_keywords, key=lambda x: -x["excess"])[:10]:
            overused_section += f"  - '{ow['keyword']}': {ow['actual']}× (max: {ow['max']}, usuń {ow['excess']}×)\n"
        overused_section += "→ Zastąp nadmiarowe wystąpienia synonimami lub zaimkami.\n"
        unused_keywords_section += overused_section
    
    # ================================================================
    # 🆕 v45.1 + v50.5: SEMANTIC COVERAGE + ENTITY PLACEMENT
    # ================================================================
    coverage_section = ""
    s1_data = project_data.get("s1_data", {})
    
    # Check entity coverage (from S1 entities)
    entity_seo = s1_data.get("entity_seo", {})
    all_entities = [e.get("name", "") for e in entity_seo.get("entities", []) if e.get("name")]
    article_lower = full_text.lower() if full_text else ""
    
    missing_entities = [e for e in all_entities if e.lower() not in article_lower]
    covered_entities = [e for e in all_entities if e.lower() in article_lower]
    entity_coverage_pct = (len(covered_entities) / len(all_entities) * 100) if all_entities else 100
    
    # v50.5: Concept entities (topical generator) — these are the KEY entities
    concept_entities = []
    for tc in entity_seo.get("topical_coverage", []):
        if isinstance(tc, dict):
            name = tc.get("entity", tc.get("name", ""))
            if name:
                concept_entities.append(name)
    
    # v50.5: Must-cover concepts from project data
    must_cover = project_data.get("must_cover_concepts", [])
    if must_cover:
        must_names = []
        for mc in must_cover:
            if isinstance(mc, dict):
                must_names.append(mc.get("text", mc.get("name", "")))
            elif isinstance(mc, str):
                must_names.append(mc)
        if must_names:
            concept_entities = must_names
    
    # Check which concept entities are missing from article
    missing_concepts = [e for e in concept_entities if e.lower() not in article_lower]
    
    # Check content gap coverage
    content_gaps = s1_data.get("content_gaps", {})
    gap_topics = []
    for gap in content_gaps.get("gaps", []):
        if isinstance(gap, dict):
            gap_topics.append(gap.get("topic", ""))
        elif isinstance(gap, str):
            gap_topics.append(gap)
    
    uncovered_gaps = []
    for gtopic in gap_topics:
        topic_words = [w for w in gtopic.lower().split() if len(w) > 3]
        if topic_words:
            found = sum(1 for w in topic_words if w in article_lower)
            if found / len(topic_words) < 0.4:
                uncovered_gaps.append(gtopic)
    
    if missing_entities or uncovered_gaps or missing_concepts or concept_entities:
        coverage_section = "\n=== 📊 POKRYCIE ENCJI I TEMATÓW ===\n"
        
        if concept_entities:
            coverage_section += f"\nENCJE TEMATYCZNE (MUST — muszą być w artykule):\n"
            for ce in concept_entities[:12]:
                present = "✅" if ce.lower() in article_lower else "❌ BRAK"
                coverage_section += f"  {present} {ce}\n"
            coverage_section += (
                "\n→ Encje oznaczone ❌ MUSZĄ pojawić się w tekście.\n"
                "→ Wpleć je naturalnie w istniejące zdania — NIE twórz nowych akapitów.\n"
                "→ Encje to POJĘCIA do opisania, NIE źródła do cytowania.\n"
                "→ ❌ NIE pisz: 'według [encji]...', '[encja] podaje...'\n"
                "→ ✅ PISZ: 'W kontekście [encji] warto zauważyć, że...'\n"
            )
        
        if missing_entities and not concept_entities:
            coverage_section += f"\nEncje z S1: {len(covered_entities)}/{len(all_entities)} pokryte ({entity_coverage_pct:.0f}%)\n"
            if missing_entities[:10]:
                coverage_section += f"BRAKUJĄCE ENCJE: {', '.join(missing_entities[:10])}\n"
        
        if uncovered_gaps[:5]:
            coverage_section += f"\nNIEPOKRYTE TEMATY (Information Gain): {', '.join(uncovered_gaps[:5])}\n"
            coverage_section += "→ Dodanie tych tematów daje przewagę nad konkurencją.\n"
    
    # 🆕 v44.5: YMYL context for editorial review
    detected_category = project_data.get("detected_category", "inne")
    ymyl_section = ""
    if detected_category == "prawo":
        ymyl_section = """
=== ⚖️ ARTYKUŁ PRAWNY (YMYL) ===
Ten artykuł dotyczy tematyki prawnej. Sprawdź SZCZEGÓLNIE:
1. Czy artykuł zawiera cytowania przepisów (art. X § Y ustawy Z)?
2. Czy są odwołania do orzeczeń sądowych (sygn. akt)?
3. Czy NIE MA wymyślonych sygnatur, dat ani przepisów?
4. Czy na końcu jest zastrzeżenie prawne (disclaimer)?
5. Jeśli brakuje cytowań — DODAJ je w sekcji "luki_tresciowe"!
"""
    elif detected_category == "medycyna":
        ymyl_section = """
=== 🏥 ARTYKUŁ MEDYCZNY (YMYL) ===
Ten artykuł dotyczy tematyki medycznej/zdrowotnej. Sprawdź SZCZEGÓLNIE:
1. Czy artykuł zawiera odwołania do publikacji naukowych?
2. Czy cytowane źródła istnieją (PMID, autorzy, rok)?
3. Czy NIE MA wymyślonych statystyk, badań ani leków?
4. Czy na końcu jest zastrzeżenie medyczne (disclaimer)?
5. Czy tekst nie zawiera niebezpiecznych rad zdrowotnych?
6. Jeśli brakuje cytowań — DODAJ je w sekcji "luki_tresciowe"!
"""

    try:
        analysis = None
        corrected_article = ""
        ai_model = "unknown"
        applied_changes = []
        failed_changes = []
        
        if claude_client:
            # ============================================================
            # WYWOŁANIE 1: RECENZJA REDAKTORSKA (v41.0)
            # ============================================================
            analysis_prompt = f"""Jesteś redaktorem naczelnym specjalistycznych serwisów branżowych.
Twoim standardem jest jakość redakcyjna mediów eksperckich — nie bloga, nie portalu generycznego.
Nie jesteś copywriterem. Oceniasz tekst jak senior editor, który odrzuca słabe materiały.

Dostajesz artykuł pt. "{topic}" ({word_count} słów). Napisz PEŁNĄ RECENZJĘ REDAKTORSKĄ.
{unused_keywords_section}
{ymyl_section}
{coverage_section}

=== NADRZĘDNA ZASADA ===
🔴 POPRAWIAJ tekst, NIE PRZEPISUJ go. Zachowaj oryginalny styl, ton i strukturę.
Każda Twoja sugestia powinna prowadzić do PUNKTOWEJ POPRAWY, nie do przepisania akapitu.

=== OBOWIĄZKOWE KOREKTY JĘZYKOWE ===
1. 🔡 WIELKA LITERA: Każde zdanie musi zaczynać się wielką literą.
2. 🔁 POWTÓRZENIA POCZĄTKÓW AKAPITÓW: Zmień jeśli akapit zaczyna się od frazy głównego słowa kluczowego.
3. 🔂 IDENTYCZNE ZDANIA: Usuń lub przepisz zdania powtórzone słowo w słowo 2+ razy.
4. 🎯 ODMIANA: Popraw encje/frazy w złym przypadku gramatycznym.

=== CZEGO SZUKAĆ — STYL I JAKOŚĆ ===

🚫 DRAMATYZATORY — znajdź i oznacz zdania będące pustymi "pointami" lub "myślami":
  Wzorzec: krótkie zdanie (1-5 słów) bez konkretnej informacji, służące "efektowi".
  Przykłady do wyeliminowania:
    ❌ "Granice są sztywne." / "Sąd patrzy. I słucha." / "I protokół."
    ❌ "To nie jest sprawa na skróty." / "Liczy się uzasadnienie."
    ❌ "W tle zostaje pytanie." / "Prawo jest bezwzględne."
  Reguła: krótkie zdanie musi nieść konkretną informację (liczbę, fakt, datę).
  ✅ OK: "Zakaz trwa od 3 do 15 lat." / "Próg to 0,5‰."

🚫 ANTY-FILLER — zdania bez informacji:
  - Truizmy: "Jazda po alkoholu to poważne przestępstwo." (oczywistość)
  - Puste przejścia: "To prowadzi do kolejnego aspektu."
  - Zapowiedzi: "Kolejna część artykułu wyjaśnia..."
  - Banały: "Warto zauważyć, że temat jest ważny."

🚫 FRAZY AI — lista zakazanych klisz:
  "warto zauważyć", "należy podkreślić", "co istotne", "kluczowe jest",
  "nie ulega wątpliwości", "warto pamiętać", "kluczowym aspektem",
  "w dzisiejszych czasach", "podsumowując", "jak wspomniano wcześniej"

🚫 ANTY-HALUCYNACJA — wymyślone dane:
  - Wymyślone statystyki, rozporządzenia, daty, ceny
  - Nieistniejące wyroki sądowe lub przepisy
  → Zaznacz jako HALUCYNACJA w errors_to_fix

🚫 WYROKI SĄDOWE (tylko dla YMYL prawo) — weryfikacja sygnatur:
  Sygnatura zdradza typ sprawy:
  • II K, III K, AKa, AKo = KARNE → pasuje do art. KK, KW
  • I C, II C, ACa, ACo = CYWILNE → pasuje do art. KC, KRO
  ❌ Wyrok cywilny (I C, II C) w artykule karnym = błąd merytoryczny
  ❌ Wyrok karny (II K) w artykule cywilnym = błąd merytoryczny
  → Zaznacz jako BŁĄD_WYROKU w errors_to_fix z sugestią usunięcia

🚫 ENCJE JAKO ŹRÓDŁA:
  - "Wikipedia podaje..." — max 1× w artykule
  - "Według [nazwy encji]..." — encje to pojęcia, nie źródła
  → Zamień na bezpośrednie stwierdzenie faktu

📊 ENCJE TEMATYCZNE: Wpleć brakujące encje z sekcji POKRYCIE ENCJI
  w istniejące zdania — NIE twórz nowych akapitów.

=== STRUKTURA RECENZJI (odpowiedz TYLKO JSON) ===

{{
  "overall_score": <0-10>,
  "scores": {{
    "merytoryka": <0-10>,
    "struktura": <0-10>,
    "styl": <0-10>,
    "seo": <0-10>
  }},

  "editorial_feedback": {{

    "recenzja_ogolna": "<3-5 zdań: ogólna ocena. Co jest mocne? Co wymaga pracy? Jaki ton i czy jest odpowiedni? Czy artykuł odpowiada na Search Intent?>",

    "merytoryka": [
      {{
        "sekcja": "<H2/H3 którego dotyczy>",
        "uwaga": "<co jest nie tak merytorycznie>",
        "cytat": "<fragment tekstu — min 10 słów>",
        "sugestia": "<jak poprawić — konkretnie>"
      }}
    ],

    "styl_i_jezyk": [
      {{
        "problem": "dramatyzator|powtórzenie|niezręczność|fraza_AI|kolokacja|strona_bierna|filler|truizm",
        "cytat": "<fragment z tekstu>",
        "sugestia": "<jak poprawić>"
      }}
    ],

    "struktura_i_narracja": [
      {{
        "uwaga": "<problem ze strukturą>",
        "gdzie": "<między którymi sekcjami>",
        "sugestia": "<co zmienić>"
      }}
    ],

    "luki_tresciowe": [
      {{
        "brakujacy_temat": "<czego czytelnik szuka, a tekst nie pokrywa>",
        "gdzie_dodac": "<w której sekcji>",
        "sugestia": "<2-3 zdania co konkretnie dopisać>"
      }}
    ],

    "halucynacje": [
      {{
        "cytat": "<fragment z wymyśloną statystyką/datą/wyrokiem>",
        "dlaczego_falsz": "<krótkie wyjaśnienie>"
      }}
    ],

    "brakujace_encje": [
      {{
        "encja": "<nazwa brakującej encji>",
        "gdzie_wplesc": "<w którym zdaniu/akapicie>",
        "jak": "<konkretna sugestia wplecenia>"
      }}
    ]
  }},

  "errors_to_fix": [
    {{
      "type": "DRAMATYZATOR|HALUCYNACJA|FILLER|FRAZA_AI|ENCJA_JAKO_ZRODLO|BŁĄD_WYROKU|KOLOKACJA|STYL|BRAK_ENCJI",
      "priority": <1-3, gdzie 1=krytyczne>,
      "original": "<cytat z tekstu — min 10 słów>",
      "replacement": "<poprawka — zachowaj długość oryginału>",
      "action": "POPRAW|USUŃ|WPLEĆ_ENCJĘ|USUŃ_WYROK"
    }}
  ],

  "keywords_to_add": {json.dumps(unused_basic + unused_extended, ensure_ascii=False)},

  "summary": "<2-3 zdania: najważniejsze co trzeba zrobić>"
}}

=== WSKAZÓWKI ===
- DRAMATYZATORY to priorytet 1 — szukaj w każdym akapicie
- W "styl_i_jezyk" szukaj: dramatyzatory, filler, frazy AI, klisze
- W "halucynacje": konkretne liczby, daty, nazwy badań, wyroki — czy brzmią wiarygodnie?
- W "brakujace_encje": encje z sekcji POKRYCIE ENCJI oznaczone ❌
- Każda poprawka w errors_to_fix musi mieć DOKŁADNY cytat (min 10 słów)
- REPLACEMENT nie może być krótszy niż ORIGINAL
- NIE przepisuj całych akapitów — poprawiaj punktowo

=== ARTYKUŁ ({word_count} słów) ===

{full_text}"""

            print(f"[EDITORIAL_REVIEW] ========== CALL 1: RECENZJA REDAKTORSKA ==========")
            
            analysis_text, ai_model = _claude_call(analysis_prompt, max_tokens=6000, label="CALL 1")
            
            # Parsuj JSON analizy
            analysis = _robust_json_parse(analysis_text)
            if analysis and analysis.get("overall_score") is not None:
                print(f"[EDITORIAL_REVIEW] ✅ Analysis parsed: score={analysis.get('overall_score')}, errors={len(analysis.get('errors_to_fix',[]))}")
            else:
                print(f"[EDITORIAL_REVIEW] ⚠️ Analysis parse failed — raw: {analysis_text[:200]}")
                analysis = {"overall_score": 5, "summary": "Analiza nie sparsowana", "errors_to_fix": [], "raw": analysis_text[:500]}
            
            # ============================================================
            # WYWOŁANIE 2: DIFF-BASED CORRECTION (v29.0)
            # ============================================================
            errors_list = analysis.get("errors_to_fix", []) if analysis else []
            keywords_to_add = analysis.get("keywords_to_add", []) if analysis else []
            
            diff_prompt = f"""Jesteś redaktorem naczelnym. Poniżej artykuł do korekty, a następnie lista błędów do poprawienia.

=== ARTYKUŁ DO KOREKTY: "{topic}" ({word_count} słów) ===

{full_text}

=== KONIEC ARTYKUŁU ===

Przeanalizuj artykuł powyżej i zwróć TYLKO ZMIANY w formacie diff.

🔴 NADRZĘDNA ZASADA: POPRAWIAJ, NIE PRZEPISUJ.
Zachowaj oryginalny styl, ton i strukturę. Zmieniaj TYLKO to co błędne.

⛔ KRYTYCZNE ZASADY:
- MAX 15 zmian (tylko najważniejsze)
- NIE przepisuj całych akapitów — to KOREKTA, nie rewrite
- Cytat w ZNAJDŹ: DOKŁADNY (min 10 słów, copy-paste z artykułu powyżej)
- ZAMIEŃ: PODOBNA długość do ZNAJDŹ (±20%)
- Wyjątek: halucynacje i błędne wyroki → USUŃ bez zastępnika

=== PRIORYTET ZMIAN ===
1. 🔴 DRAMATYZATORY — krótkie zdania-pointy bez informacji → zastąp zdaniem z konkretną informacją
2. 🔴 HALUCYNACJE — wymyślone statystyki, rozporządzenia, daty → USUŃ lub zastąp pewnym faktem
3. 🔴 BŁĘDY WYROKÓW — wyrok cywilny (I C) w artykule karnym lub odwrotnie → USUŃ_WYROK
4. 🔴 FILLER/TRUIZMY — zdania bez informacji → zastąp treścią merytoryczną
5. 🔴 ENCJA JAKO ŹRÓDŁO — "Wikipedia podaje...", "Według [encji]..." → bezpośrednie stwierdzenie
6. 🟡 BRAKUJĄCE ENCJE — wpleć w istniejące zdania
7. 🟡 BRAKUJĄCE FRAZY SEO — wpleć naturalnie
8. 🟢 STYL — frazy AI, złe kolokacje, powtórzenia

=== BŁĘDY DO POPRAWY (z analizy redaktorskiej) ===
{json.dumps(errors_list[:12], ensure_ascii=False, indent=2) if errors_list else "Brak krytycznych błędów."}

=== FRAZY DO WPLECENIA ===
{', '.join(keywords_to_add[:10]) if keywords_to_add else "Wszystkie frazy są w tekście."}

{coverage_section}

=== FORMAT ODPOWIEDZI ===

Każda zmiana w dokładnie tym formacie:

[ZMIANA 1]
ZNAJDŹ: "dokładny cytat z artykułu powyżej (min 10 słów)"
ZAMIEŃ: "poprawiona wersja o podobnej długości"
POWÓD: typ błędu (max 5 słów)

Przykład poprawnej zmiany:
[ZMIANA 1]
ZNAJDŹ: "Granice są sztywne. Sąd nie ma tu wiele do powiedzenia i musi działać zgodnie z przepisami."
ZAMIEŃ: "Sąd nie ma tu uznaniowości — zakaz prowadzenia pojazdów jest obligatoryjny przy każdym wyroku skazującym z art. 178a KK."
POWÓD: dramatyzator → konkretna informacja prawna

Przykład usunięcia błędnego wyroku:
[ZMIANA 2]
ZNAJDŹ: "Potwierdza to wyrok Sądu Okręgowego w Słupsku z dnia 15 marca 2021 r. (sygn. I C 245/21), który wskazał na konieczność"
ZAMIEŃ: [USUŃ — wyrok cywilny (I C) w artykule karnym]
POWÓD: błędny wyrok — cywilny w artykule karnym

(kontynuuj do max 15 zmian)"""

            print(f"[EDITORIAL_REVIEW] ========== CALL 2: DIFF-BASED CORRECTION ==========")
            
            diff_response, _ = _claude_call(diff_prompt, max_tokens=8000, label="CALL 2")
            
            # Parsuj diffy
            changes = parse_diff_response(diff_response)
            print(f"[EDITORIAL_REVIEW] 📝 Parsed {len(changes)} changes from Claude response")
            
            # Aplikuj diffy
            if changes:
                corrected_article, applied_changes, failed_changes = apply_diffs(full_text, changes)
                print(f"[EDITORIAL_REVIEW] ✅ Applied: {len(applied_changes)}, ❌ Failed: {len(failed_changes)}")
            else:
                corrected_article = full_text
                print(f"[EDITORIAL_REVIEW] ℹ️ No changes parsed, using original text")
            
        elif genai:
            # Gemini fallback (stary sposób - przepisuje całość)
            print(f"[EDITORIAL_REVIEW] Using Gemini (fallback) for project {project_id}")
            model = genai.GenerativeModel("gemini-2.0-flash")
            
            old_prompt = f"""Jesteś REDAKTOREM. Popraw artykuł i zwróć JSON:
{{"analysis": {{"overall_score": <0-10>, "summary": "<podsumowanie>"}}, "corrected_article": "<cały poprawiony tekst>"}}

Artykuł ({word_count} słów):
{full_text[:10000]}"""
            
            response = model.generate_content(old_prompt)
            review_text = response.text.strip()
            ai_model = "gemini"
            
            try:
                clean = review_text
                if "```json" in clean:
                    clean = re.sub(r'```json\s*', '', clean)
                    clean = re.sub(r'```\s*$', '', clean)
                data = _robust_json_parse(clean)
                if not data:
                    data = {}
                analysis = data.get("analysis", {'overall_score': 5, 'summary': 'parse incomplete'})
                corrected_article = data.get("corrected_article", "") or ""
                if not corrected_article:
                    corrected_article = full_text
            except:
                analysis = {"overall_score": 5, "summary": "Gemini parse error"}
                corrected_article = full_text
        else:
            return jsonify({
                "error": "No AI API configured",
                "hint": "Set ANTHROPIC_API_KEY or GEMINI_API_KEY"
            }), 500
        
        # Fallback jeśli brak corrected_article
        if not corrected_article or len(corrected_article) < 100:
            print(f"[EDITORIAL_REVIEW] ⚠️ Using original text as fallback")
            corrected_article = full_text
        
        # ================================================================
        # v41.0: WORD COUNT GUARD — artykuł nie powinien się skrócić
        # ================================================================
        corrected_wc = len(corrected_article.split())
        word_count_shrinkage = False
        shrinkage_pct = 0
        
        if corrected_wc < word_count and corrected_article != full_text:
            shrinkage_pct = round((1 - corrected_wc / word_count) * 100, 1)
            print(f"[EDITORIAL_REVIEW] ⚠️ WORD COUNT SHRINKAGE: {word_count} → {corrected_wc} (-{shrinkage_pct}%)")

            if shrinkage_pct > 5:
                # Check surgery_applied from project_data Firestore
                batches = project_data.get('batches', []) or []
                surgery_applied = False
                if batches:
                    last_moe = (batches[-1].get('moe_validation') or {})
                    surgery_applied = last_moe.get('surgery_applied', False)

                if surgery_applied and shrinkage_pct <= 15:
                    print(f"[EDITORIAL] Surgery shrinkage {shrinkage_pct}% accepted")
                else:
                    word_count_shrinkage = True
                    print(f"[EDITORIAL] ROLLBACK: {shrinkage_pct}%")
                    corrected_article = full_text
        
        # ================================================================
        # v55.1: SENTENCE LENGTH GUARD — editorial nie powinien robić zdań >20 słów
        # ================================================================
        if corrected_article != full_text:
            import re as _re
            _sentences = [s.strip() for s in _re.split(r'[.!?]+', corrected_article) if len(s.strip().split()) >= 3]
            if _sentences:
                avg_sent_len = sum(len(s.split()) for s in _sentences) / len(_sentences)
                if avg_sent_len > 22:
                    # Check if original was also long — only rollback if editorial made it worse
                    _orig_sentences = [s.strip() for s in _re.split(r'[.!?]+', full_text) if len(s.strip().split()) >= 3]
                    orig_avg = sum(len(s.split()) for s in _orig_sentences) / len(_orig_sentences) if _orig_sentences else 15
                    if avg_sent_len > orig_avg + 3:
                        print(f"[EDITORIAL_REVIEW] ⚠️ SENTENCE LENGTH ROLLBACK: avg {avg_sent_len:.1f} > orig {orig_avg:.1f}+3 — reverting")
                        corrected_article = full_text
                        rollback_reason = f"Sentence length degradation: {avg_sent_len:.1f} > {orig_avg:.1f}+3"
                        rollback_triggered = True

        # ================================================================
        # v53: GRAMMAR CORRECTION — BEFORE burstiness check
        # Moved here so grammar fixes are not lost to rollback.
        # ================================================================
        grammar_stats = {"fixes": 0, "removed": []}
        if corrected_article != full_text:
            try:
                from grammar_checker import full_correction
                grammar_result = full_correction(corrected_article)
                corrected_article = grammar_result["corrected"]
                grammar_stats = {
                    "fixes": grammar_result["grammar_fixes"],
                    "removed": grammar_result["phrases_removed"],
                    "backend": grammar_result.get("backend", "unknown")
                }
                if grammar_stats["fixes"] > 0 or grammar_stats["removed"]:
                    print(f"[EDITORIAL_REVIEW] ✅ Grammar (pre-rollback): {grammar_stats['fixes']} fixes, {len(grammar_stats['removed'])} phrases removed")
            except ImportError:
                print(f"[EDITORIAL_REVIEW] ⚠️ grammar_checker not available, skipping grammar correction")
            except Exception as e:
                print(f"[EDITORIAL_REVIEW] ⚠️ Grammar correction error: {e}")

        # ================================================================
        # v53: BURSTINESS CHECK + GRANULAR ROLLBACK
        # Instead of rolling back ALL changes, test each individually
        # and keep those that don't harm burstiness.
        # ================================================================
        burstiness_after = None
        rollback_triggered = False
        rollback_reason = None
        granular_kept = 0
        granular_reverted = 0

        if BURSTINESS_CHECK_OK and burstiness_before and corrected_article != full_text:
            try:
                burstiness_after = calculate_burstiness(corrected_article)
                cv_before = burstiness_before.get("cv", 0)
                cv_after = burstiness_after.get("cv", 0)

                print(f"[EDITORIAL_REVIEW] 📊 Burstiness PO: CV={cv_after}")

                cv_drop = cv_before - cv_after
                needs_granular = False

                if cv_drop > 0.2 and cv_after < 0.35:
                    needs_granular = True
                    rollback_reason = f"CV spadło o {cv_drop:.2f} (z {cv_before:.2f} do {cv_after:.2f}) — uruchamiam granularny rollback"
                    print(f"[EDITORIAL_REVIEW] ⚠️ {rollback_reason}")

                elif cv_after < 0.25 and cv_before >= 0.3:
                    needs_granular = True
                    rollback_reason = f"CV spadło poniżej progu AI ({cv_after:.2f} < 0.25) — uruchamiam granularny rollback"
                    print(f"[EDITORIAL_REVIEW] ⚠️ {rollback_reason}")

                elif cv_drop > 0.1:
                    print(f"[EDITORIAL_REVIEW] ⚠️ CV drop {cv_drop:.2f} (z {cv_before:.2f} do {cv_after:.2f}) — akceptowalne (cv_after={cv_after:.2f} >= 0.35)")

                # ── GRANULAR ROLLBACK ──
                # Apply changes one by one, keep those that don't break CV
                if needs_granular and applied_changes:
                    print(f"[EDITORIAL_REVIEW] 🔬 Granular rollback: testing {len(applied_changes)} changes individually...")

                    # Start from original text
                    granular_text = full_text
                    kept_changes = []
                    cv_min_threshold = 0.30  # absolute minimum CV we accept

                    for i, change in enumerate(applied_changes):
                        # Try applying this single change
                        test_text = granular_text
                        find_text = change.get("find", "")
                        replace_text = change.get("replace", "")

                        if not find_text or not replace_text:
                            continue

                        # Try exact match first, then quote-normalized
                        if find_text in test_text:
                            test_text = test_text.replace(find_text, replace_text, 1)
                        elif _normalize_quotes(find_text) in _normalize_quotes(test_text):
                            fn = _normalize_quotes(find_text)
                            tn = _normalize_quotes(test_text)
                            idx = tn.index(fn)
                            test_text = test_text[:idx] + replace_text + test_text[idx + len(fn):]
                        else:
                            # Can't apply this change to original — skip
                            granular_reverted += 1
                            continue

                        # Check CV after this change
                        test_burst = calculate_burstiness(test_text)
                        test_cv = test_burst.get("cv", 0)

                        if test_cv >= cv_min_threshold:
                            # Safe — keep this change
                            granular_text = test_text
                            kept_changes.append(change)
                            granular_kept += 1
                        else:
                            # This change breaks CV — revert it
                            granular_reverted += 1
                            print(f"[EDITORIAL_REVIEW] 🔬 Change {i+1} reverted (CV would drop to {test_cv:.2f}): {find_text[:50]}...")

                    corrected_article = granular_text

                    # Recalculate final burstiness
                    burstiness_after = calculate_burstiness(corrected_article)
                    final_cv = burstiness_after.get("cv", 0)

                    if granular_kept > 0:
                        rollback_triggered = False  # Partial success — not a full rollback
                        rollback_reason = f"Granular: zachowano {granular_kept}/{granular_kept + granular_reverted} zmian (CV: {cv_before:.2f} → {final_cv:.2f})"
                        print(f"[EDITORIAL_REVIEW] ✅ {rollback_reason}")
                    else:
                        rollback_triggered = True
                        rollback_reason = f"Granular rollback: żadna zmiana nie przeszła progu CV (CV: {cv_before:.2f} → {cv_after:.2f})"
                        corrected_article = full_text
                        print(f"[EDITORIAL_REVIEW] ⚠️ {rollback_reason}")

                elif needs_granular and not applied_changes:
                    # No individual changes to test — full rollback
                    rollback_triggered = True
                    corrected_article = full_text
                    print(f"[EDITORIAL_REVIEW] ⚠️ Full rollback (no individual changes to test)")

            except Exception as e:
                print(f"[EDITORIAL_REVIEW] ⚠️ Burstiness after check failed: {e}")
        
        # v41.0: Rollback z powodu skrócenia artykułu
        if word_count_shrinkage and not rollback_triggered:
            rollback_triggered = True
            rollback_reason = f"Artykuł skrócił się o {shrinkage_pct}% ({word_count} → {corrected_wc} słów) — editorial review powinien ROZBUDOWYWAĆ, nie skracać"
        
        # Statystyki diff
        diff_stats = calculate_diff_stats(full_text, corrected_article)

        # v53: grammar_stats already computed above (before burstiness check)
        corrected_word_count = len(corrected_article.split())
        
        # Zapisz w Firestore
        db.collection("seo_projects").document(project_id).update({
            "editorial_review": {
                "analysis": analysis,
                "corrected_article": corrected_article,
                "original_article": full_text,  # v29.0: Zawsze zachowuj oryginał!
                "original_word_count": word_count,
                "corrected_word_count": corrected_word_count,
                "unused_keywords": {
                    "basic": unused_basic,
                    "extended": unused_extended
                },
                # v29.0: Nowe pola
                "applied_changes": applied_changes[:20],
                "failed_changes": failed_changes[:10],
                "burstiness_before": burstiness_before,
                "burstiness_after": burstiness_after,
                "rollback_triggered": rollback_triggered,
                "rollback_reason": rollback_reason,
                "diff_stats": diff_stats,
                "grammar_correction": grammar_stats,
                "ai_model": ai_model,
                "version": "v41.0-expand",
                "created_at": firestore.SERVER_TIMESTAMP
            }
        })
        
        # v27.1: RESCAN KEYWORDS
        rescan_result = {"rescanned": False, "reason": "no_corrected_article"}
        if not rollback_triggered and corrected_article and len(corrected_article.strip()) > 50:
            print(f"[EDITORIAL_REVIEW] Running rescan_keywords for project {project_id}...")
            rescan_result = rescan_keywords_after_editorial(project_id, corrected_article)
            print(f"[EDITORIAL_REVIEW] Rescan result: {rescan_result.get('changes_count', 0)} changes")
        
        return jsonify({
            "status": "REVIEWED",
            "version": "v41.0-DIFF-EXPAND",
            "overall_score": analysis.get("overall_score") if analysis else None,
            "scores": analysis.get("scores", {}) if analysis else {},
            "summary": analysis.get("summary", "") if analysis else "",
            
            # 🆕 v41.0: PEŁNA RECENZJA REDAKTORSKA
            "editorial_feedback": analysis.get("editorial_feedback", {}) if analysis else {},
            
            # v29.0: DIFF OUTPUT
            "diff_result": {
                "total_changes_parsed": len(applied_changes) + len(failed_changes),
                "applied": len(applied_changes),
                "failed": len(failed_changes),
                "applied_changes": applied_changes[:10],
                "failed_changes": failed_changes[:5]
            },
            "diff_stats": diff_stats,
            
            # v29.0: BURSTINESS VALIDATION
            "burstiness_check": {
                "enabled": BURSTINESS_CHECK_OK,
                "before": {
                    "cv": burstiness_before.get("cv") if burstiness_before else None,
                    "value": burstiness_before.get("value") if burstiness_before else None,
                    "status": burstiness_before.get("status") if burstiness_before else None
                },
                "after": {
                    "cv": burstiness_after.get("cv") if burstiness_after else None,
                    "value": burstiness_after.get("value") if burstiness_after else None,
                    "status": burstiness_after.get("status") if burstiness_after else None
                },
                "cv_change": round((burstiness_after.get("cv", 0) - burstiness_before.get("cv", 0)), 3) if (burstiness_before and burstiness_after) else None
            },
            
            # v41.0: WORD COUNT GUARD
            "word_count_guard": {
                "original": word_count,
                "corrected": corrected_word_count,
                "shrinkage_pct": shrinkage_pct,
                "rollback_triggered": word_count_shrinkage,
                "policy": "Editorial review powinien ROZBUDOWYWAĆ treść, nie skracać"
            },
            
            # v29.0: ROLLBACK INFO
            "rollback": {
                "triggered": rollback_triggered,
                "reason": rollback_reason,
                "original_preserved": True,
                "granular_kept": granular_kept,
                "granular_reverted": granular_reverted,
            },
            
            "corrected_article": corrected_article,
            "word_count": {
                "original": word_count,
                "corrected": corrected_word_count
            },
            "unused_keywords_input": {
                "basic": unused_basic,
                "extended": unused_extended
            },
            "ai_model": ai_model,
            "grammar_correction": grammar_stats,
            "rescan_result": rescan_result,
            
            # 🆕 v44.5: YMYL validation
            "ymyl_check": {
                "detected_category": detected_category,
                "is_ymyl": detected_category in ("prawo", "medycyna"),
                "ymyl_reviewed": detected_category in ("prawo", "medycyna")
            } if detected_category in ("prawo", "medycyna") else None
        }), 200
        
    except Exception as e:
        import traceback
        tb = traceback.format_exc()
        print(f"[EDITORIAL_REVIEW] ❌ Error: {e}")
        print(f"[EDITORIAL_REVIEW] Traceback:\n{tb}")
        return jsonify({
            "error": f"Review failed: {str(e)}",
            "error_type": type(e).__name__,
            "word_count": word_count if 'word_count' in dir() else 0,
            "debug": debug_info if 'debug_info' in dir() else {},
            "traceback_hint": tb.strip().split("\n")[-1] if tb else ""
        }), 500


# ================================================================
# v29.0: MANUAL ROLLBACK ENDPOINT
# ================================================================
@export_routes.post("/api/project/<project_id>/editorial_review/rollback")
def editorial_rollback(project_id):
    """
    v29.0: Ręczny rollback do oryginalnego tekstu.
    Użyj gdy automatyczny rollback nie zadziałał lub chcesz przywrócić oryginał.
    """
    db = firestore.client()
    doc = db.collection("seo_projects").document(project_id).get()
    
    if not doc.exists:
        return jsonify({"error": "Project not found"}), 404
    
    project_data = doc.to_dict()
    editorial = project_data.get("editorial_review", {})
    
    original = editorial.get("original_article")
    if not original:
        return jsonify({
            "error": "No original article saved",
            "hint": "Original article is only available after editorial_review v29.0+"
        }), 400
    
    current = editorial.get("corrected_article", "")
    
    if original == current:
        return jsonify({
            "status": "NO_CHANGE",
            "message": "Tekst już jest oryginalny (poprzedni rollback lub brak zmian)"
        }), 200
    
    # Przywróć oryginał
    db.collection("seo_projects").document(project_id).update({
        "editorial_review.corrected_article": original,
        "editorial_review.manual_rollback": True,
        "editorial_review.rollback_at": firestore.SERVER_TIMESTAMP,
        "editorial_review.pre_rollback_article": current
    })
    
    return jsonify({
        "status": "ROLLED_BACK",
        "message": "Przywrócono oryginalny tekst",
        "word_count": {
            "original": len(original.split()),
            "was_corrected": len(current.split())
        }
    }), 200


# ================================================================
# EXPORT STATUS
# ================================================================
@export_routes.get("/api/project/<project_id>/export_status")
def export_status(project_id):
    """Sprawdza status projektu przed eksportem."""
    db = firestore.client()
    doc = db.collection("seo_projects").document(project_id).get()
    
    if not doc.exists:
        return jsonify({"error": "Project not found"}), 404
    
    project_data = doc.to_dict()
    
    full_text = None
    debug_info = {"source": "none"}
    batches_count = 0
    
    if _safe_get_full_article(project_data):
        full_text = _safe_get_full_article(project_data)
        debug_info = {"source": "full_article"}
    else:
        batches = project_data.get("batches", [])
        batches_count = len(batches)
        if batches:
            full_text, debug_info = extract_text_from_batches(batches)
            debug_info["source"] = "batches"
    
    word_count = len(full_text.split()) if full_text else 0
    
    editorial = project_data.get("editorial_review", {})
    
    return jsonify({
        "project_id": project_id,
        "topic": project_data.get("topic", ""),
        "has_content": bool(full_text),
        "word_count": word_count,
        "batches_count": batches_count,
        "content_source": debug_info.get("source", "none"),
        "editorial_review": {
            "done": bool(editorial),
            "score": editorial.get("analysis", {}).get("overall_score") if editorial else None,
            "version": editorial.get("version", "unknown"),
            "rollback_triggered": editorial.get("rollback_triggered", False),
            "burstiness_cv_before": editorial.get("burstiness_before", {}).get("cv") if editorial.get("burstiness_before") else None,
            "burstiness_cv_after": editorial.get("burstiness_after", {}).get("cv") if editorial.get("burstiness_after") else None
        },
        "available_exports": {
            "docx": f"/api/project/{project_id}/export/docx",
            "html": f"/api/project/{project_id}/export/html",
            "txt": f"/api/project/{project_id}/export/txt"
        },
        "next_steps": {
            "editorial_review": f"POST /api/project/{project_id}/editorial_review",
            "export_docx": f"GET /api/project/{project_id}/export/docx",
            "export_html": f"GET /api/project/{project_id}/export/html"
        }
    }), 200


# ================================================================
# EXPORT DOCX
# ================================================================
@export_routes.get("/api/project/<project_id>/export/docx")
def export_docx(project_id):
    """Eksportuje artykuł do formatu DOCX."""
    db = firestore.client()
    doc = db.collection("seo_projects").document(project_id).get()
    
    if not doc.exists:
        return jsonify({"error": "Project not found"}), 404
    
    project_data = doc.to_dict()
    
    # v32.4: Priorytet źródeł treści
    # 1. editorial_review.corrected_article (po Claude review)
    # 2. full_article.content (zapisany przez save_full_article)
    # 3. batches (surowe batche)
    editorial = project_data.get("editorial_review", {})
    corrected = editorial.get("corrected_article", "")
    
    if corrected:
        full_text = corrected
        print(f"[EXPORT_DOCX] Using corrected article ({len(corrected.split())} words)")
    elif _safe_get_full_article(project_data):
        full_text = _safe_get_full_article(project_data)
        print(f"[EXPORT_DOCX] Using full_article ({len(full_text.split())} words)")
    else:
        batches = project_data.get("batches", [])
        full_text, _ = extract_text_from_batches(batches)
        print(f"[EXPORT_DOCX] Using batches ({len(full_text.split()) if full_text else 0} words)")
    
    if not full_text:
        return jsonify({"error": "No content to export"}), 400
    
    # 🆕 v44.5: YMYL disclaimer guard
    detected_category = project_data.get("detected_category", "inne")
    if detected_category == "prawo" and "zastrzeżenie" not in full_text.lower():
        full_text += (
            "\n\n---\n\n"
            "**Zastrzeżenie prawne:** Niniejszy artykuł ma charakter wyłącznie informacyjny "
            "i nie stanowi porady prawnej. W indywidualnych sprawach zalecamy konsultację "
            "z wykwalifikowanym prawnikiem."
        )
        print(f"[EXPORT_DOCX] ⚖️ Auto-dodano brakujący disclaimer prawny")
    elif detected_category == "medycyna" and "zastrzeżenie" not in full_text.lower():
        full_text += (
            "\n\n---\n\n"
            "**Zastrzeżenie medyczne:** Niniejszy artykuł ma charakter wyłącznie informacyjny "
            "i edukacyjny. Nie stanowi porady medycznej ani nie zastępuje konsultacji "
            "z lekarzem lub innym wykwalifikowanym specjalistą."
        )
        print(f"[EXPORT_DOCX] 🏥 Auto-dodano brakujący disclaimer medyczny")
    
    topic = project_data.get("topic", "Artykuł")
    elements = parse_article_structure(full_text)
    
    # Twórz dokument
    document = Document()
    
    # Tytuł
    title = document.add_heading(topic, 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Treść
    for el in elements:
        if el["type"] == "h2":
            document.add_heading(el["content"], level=2)
        elif el["type"] == "h3":
            document.add_heading(el["content"], level=3)
        elif el["type"] == "paragraph":
            p = document.add_paragraph(el["content"])
            p.style.font.size = Pt(11)
    
    # Zapisz do bufora
    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)
    
    # Bezpieczna nazwa pliku
    polish_map = str.maketrans({
        'ą': 'a', 'ć': 'c', 'ę': 'e', 'ł': 'l', 'ń': 'n',
        'ó': 'o', 'ś': 's', 'ź': 'z', 'ż': 'z',
        'Ą': 'A', 'Ć': 'C', 'Ę': 'E', 'Ł': 'L', 'Ń': 'N',
        'Ó': 'O', 'Ś': 'S', 'Ź': 'Z', 'Ż': 'Z'
    })
    safe_topic = topic.translate(polish_map)
    filename = re.sub(r'[^\w\s-]', '', safe_topic)[:50] + ".docx"
    
    from urllib.parse import quote
    filename_utf8 = quote(re.sub(r'[^\w\s-]', '', topic)[:50] + ".docx")
    
    return Response(
        buffer.getvalue(),
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={
            "Content-Disposition": f"attachment; filename={filename}; filename*=UTF-8''{filename_utf8}"
        }
    )


# ================================================================
# EXPORT HTML
# ================================================================
@export_routes.get("/api/project/<project_id>/export/html")
def export_html(project_id):
    """Eksportuje artykuł do formatu HTML."""
    db = firestore.client()
    doc = db.collection("seo_projects").document(project_id).get()
    
    if not doc.exists:
        return jsonify({"error": "Project not found"}), 404
    
    project_data = doc.to_dict()
    
    editorial = project_data.get("editorial_review", {})
    corrected = editorial.get("corrected_article", "")
    
    if corrected:
        full_text = corrected
    elif _safe_get_full_article(project_data):
        full_text = _safe_get_full_article(project_data)
    else:
        batches = project_data.get("batches", [])
        full_text, _ = extract_text_from_batches(batches)
    
    if not full_text:
        return jsonify({"error": "No content to export"}), 400
    
    # 🆕 v44.5: YMYL disclaimer guard (HTML export)
    detected_category = project_data.get("detected_category", "inne")
    if detected_category == "prawo" and "zastrzeżenie" not in full_text.lower():
        full_text += (
            "\n\n---\n\n"
            "**Zastrzeżenie prawne:** Niniejszy artykuł ma charakter wyłącznie informacyjny "
            "i nie stanowi porady prawnej. W indywidualnych sprawach zalecamy konsultację "
            "z wykwalifikowanym prawnikiem."
        )
    elif detected_category == "medycyna" and "zastrzeżenie" not in full_text.lower():
        full_text += (
            "\n\n---\n\n"
            "**Zastrzeżenie medyczne:** Niniejszy artykuł ma charakter wyłącznie informacyjny "
            "i edukacyjny. Nie stanowi porady medycznej ani nie zastępuje konsultacji "
            "z lekarzem lub innym wykwalifikowanym specjalistą."
        )
    
    topic = project_data.get("topic", "Artykuł")
    
    # Konwertuj markery na HTML
    html = full_text
    html = re.sub(r'^h2:\s*(.+)$', r'<h2>\1</h2>', html, flags=re.MULTILINE)
    html = re.sub(r'^h3:\s*(.+)$', r'<h3>\1</h3>', html, flags=re.MULTILINE)
    
    # Paragrafy
    lines = html.split('\n')
    processed = []
    for line in lines:
        line = line.strip()
        if line and not line.startswith('<h'):
            if line.startswith('- ') or line.startswith('• '):
                processed.append(f'<li>{line[2:]}</li>')
            else:
                processed.append(f'<p>{line}</p>')
        else:
            processed.append(line)
    
    html_content = f"""<!DOCTYPE html>
<html lang="pl">
<head>
    <meta charset="UTF-8">
    <title>{topic}</title>
    <style>
        body {{ font-family: Arial, sans-serif; max-width: 800px; margin: 0 auto; padding: 20px; line-height: 1.6; }}
        h1 {{ color: #333; }}
        h2 {{ color: #444; margin-top: 30px; }}
        h3 {{ color: #555; }}
        p {{ margin: 15px 0; }}
    </style>
</head>
<body>
    <h1>{topic}</h1>
    {''.join(processed)}
</body>
</html>"""
    
    polish_map = str.maketrans({
        'ą': 'a', 'ć': 'c', 'ę': 'e', 'ł': 'l', 'ń': 'n',
        'ó': 'o', 'ś': 's', 'ź': 'z', 'ż': 'z',
        'Ą': 'A', 'Ć': 'C', 'Ę': 'E', 'Ł': 'L', 'Ń': 'N',
        'Ó': 'O', 'Ś': 'S', 'Ź': 'Z', 'Ż': 'Z'
    })
    safe_topic = topic.translate(polish_map)
    filename = re.sub(r'[^\w\s-]', '', safe_topic)[:50] + ".html"
    
    from urllib.parse import quote
    filename_utf8 = quote(re.sub(r'[^\w\s-]', '', topic)[:50] + ".html")
    
    return Response(
        html_content,
        mimetype="text/html",
        headers={
            "Content-Disposition": f"attachment; filename={filename}; filename*=UTF-8''{filename_utf8}"
        }
    )


# ================================================================
# EXPORT TXT
# ================================================================
@export_routes.get("/api/project/<project_id>/export/txt")
def export_txt(project_id):
    """Eksportuje artykuł do formatu TXT."""
    db = firestore.client()
    doc = db.collection("seo_projects").document(project_id).get()
    
    if not doc.exists:
        return jsonify({"error": "Project not found"}), 404
    
    project_data = doc.to_dict()
    
    editorial = project_data.get("editorial_review", {})
    corrected = editorial.get("corrected_article", "")
    
    if corrected:
        full_text = corrected
    elif _safe_get_full_article(project_data):
        full_text = _safe_get_full_article(project_data)
    else:
        batches = project_data.get("batches", [])
        full_text, _ = extract_text_from_batches(batches)
    
    if not full_text:
        return jsonify({"error": "No content to export"}), 400
    
    topic = project_data.get("topic", "Artykuł")
    
    # Konwertuj do czytelnego tekstu
    txt = f"{topic}\n{'='*len(topic)}\n\n"
    txt += re.sub(r'^h2:\s*(.+)$', r'\n## \1\n', full_text, flags=re.MULTILINE)
    txt = re.sub(r'^h3:\s*(.+)$', r'\n### \1\n', txt, flags=re.MULTILINE)
    
    polish_map = str.maketrans({
        'ą': 'a', 'ć': 'c', 'ę': 'e', 'ł': 'l', 'ń': 'n',
        'ó': 'o', 'ś': 's', 'ź': 'z', 'ż': 'z',
        'Ą': 'A', 'Ć': 'C', 'Ę': 'E', 'Ł': 'L', 'Ń': 'N',
        'Ó': 'O', 'Ś': 'S', 'Ź': 'Z', 'Ż': 'Z'
    })
    safe_topic = topic.translate(polish_map)
    filename = re.sub(r'[^\w\s-]', '', safe_topic)[:50] + ".txt"
    
    from urllib.parse import quote
    filename_utf8 = quote(re.sub(r'[^\w\s-]', '', topic)[:50] + ".txt")
    
    return Response(
        txt,
        mimetype="text/plain; charset=utf-8",
        headers={
            "Content-Disposition": f"attachment; filename={filename}; filename*=UTF-8''{filename_utf8}"
        }
    )


# ================================================================
# ALIASY dla kompatybilności
# ================================================================
@export_routes.post("/api/project/<project_id>/gemini_review")
def gemini_review_alias(project_id):
    """Alias do editorial_review dla kompatybilności"""
    return editorial_review(project_id)


@export_routes.post("/api/project/<project_id>/claude_review")
def claude_review_alias(project_id):
    """Alias do editorial_review"""
    return editorial_review(project_id)
